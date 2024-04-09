import fitz  # 用于操作PDF文件
import docx  # 用于操作Word (.docx) 文件
import json  # 用于处理JSON数据
import jsonlines  # 用于处理逐行的JSON数据
from tqdm import tqdm  # 用于在循环操作中添加进度条
import time  # 用于处理时间相关的功能
import validators  # 用于验证URL
from bs4 import BeautifulSoup  # 用于解析HTML和XML文档
import requests  # 用于发起HTTP请求
# import openai  # 用于调用OpenAI的API
import numpy as np  # 数学库，用于高效的多维数组操作
from numpy.linalg import norm  # 用于向量和矩阵的数学运算
import os  # 用于操作系统级别的调用，比如文件路径操作
import hashlib  # 用于哈希加密功能
import tiktoken  # 指定如何将文本转换为tokens

from gpt35 import makedata, url, headers

from haystack.schema import Document
from haystack.document_stores import InMemoryDocumentStore
import rocketqa
from rocket_qa_retriever import RocketQAEmbeddingRetriever
import threading

tokenizer = tiktoken.get_encoding("cl100k_base")

model = rocketqa.load_model(model="zh_dureader_de", use_cuda=False)


def get_doc_retriever(document_store=None):
    retriever = RocketQAEmbeddingRetriever(
        model=model,
        document_store=document_store,
        use_gpu=False,
        is_faq=False,
        scale_score=False,
    )
    return retriever


doc_retriever = get_doc_retriever()  # 文档检索器，全局变量

locker = threading.Lock()

document_store = InMemoryDocumentStore()

# with open("openai_api_key.txt", 'r', encoding='utf8') as f:
#     openai.api_key = f.readlines()[0].strip()
# print("Loaded openai api key.")


# 在控制台输出中使用不同的颜色
class bcolors:
    HEADER = '\033[95m'  # 紫色
    OKBLUE = '\033[94m'  # 蓝色
    OKCYAN = '\033[96m'  # 青色
    OKGREEN = '\033[92m'  # 绿色
    WARNING = '\033[93m'  # 黄色
    FAIL = '\033[91m'  # 红色
    ENDC = '\033[0m'  # 结束颜色标记
    BOLD = '\033[1m'  # 粗体
    UNDERLINE = '\033[4m'  # 下划线


# read file

def get_text(text_path):
    # 从不同格式的文件或网页URL中提取文本内容

    # 确定输入路径的后缀名，以判断文件类型或是网址
    url = text_path
    suffix = os.path.splitext(text_path)[-1]

    # 如果是一个有效的URL，则尝试从网页上抓取文本
    if validators.url(url):
        # 设置HTTP请求的头部信息，伪装成浏览器访问，以避免一些网站的访问限制
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36", }
        response = requests.get(url, headers=headers)
        # 如果请求成功（HTTP状态码200），则解析HTML内容并提取文本
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, "html.parser")
            text = soup.get_text()
        else:
            # 如果URL无效或无法访问，抛出异常
            raise ValueError(
                f"Invalid URL! Status code {response.status_code}.")

    # 如果文件后缀名为.pdf，使用fitz库处理PDF文件
    elif suffix == ".pdf":
        full_text = ""
        num_pages = 0
        with fitz.open(text_path) as doc:
            for page in doc:
                num_pages += 1
                text = page.get_text()
                full_text += text + "\n"
        text = f"This is a {num_pages}-page document.\n" + full_text

    # 如果文件后缀名为.doc或.docx，使用docx库处理Word文档
    elif ".doc" in suffix:
        doc = docx.Document(text_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        text = '\n'.join(fullText)

    # 如果文件是一个纯文本文件（.txt）
    elif suffix == ".txt":
        with open(text_path, 'r', encoding='utf8') as f:
            lines = f.readlines()
        text = '\n'.join(lines)

    else:
        # 如果文件格式不受支持，抛出异常
        raise ValueError("Invalid document path!")

    # 最后，将文本中的多余空格合并，并返回处理后的文本
    text = " ".join(text.split())
    return text

# embedding


def get_embedding(text):
    """
    根据给定的文本生成嵌入向量(embedding)

    args:
    - text: 需要生成嵌入向量的文本，类型为str
    - model: 使用的模型名称，默认为"text-embedding-ada-002"

    returns:
    - 返回一个表示文本嵌入向量的列表
    """

    # 将输入文本中的换行符替换为单个空格，以避免在处理文本时发生潜在的格式问题
    text = text.replace("\n", " ")

    # 使用指定的模型生成文本的嵌入向量
    # 这里使用OpenAI提供的Embedding API，输入为处理后的文本和模型名称
    # ['data'][0]['embedding']返回的是嵌入向量的具体数据
    # return openai.Embedding.create(input=[text], model=model)['data'][0]['embedding']

    locker.acquire()
    embedding = doc_retriever.embed_documents([Document(content=text)])[
        0]  # 使用retriever生成嵌入向量
    locker.release()
    return embedding


def store_info(text, chunk_sz=30, max_memory=100):
    # 英语500
    # 中文40
    """
    处理文本，分块生成摘要和嵌入向量，并存储这些信息。

    args:
    - text: 需要处理的原始文本，str类型。
    - memory_path: 存储处理后信息的文件路径，str类型。
    - chunk_sz: 每个文本块的最大单词数，默认700，int类型。
    - max_memory: 允许的最大API调用次数，用于控制成本，默认100，int类型。

    raises:
    - ValueError: 如果预计的API使用量过高，则抛出错误。
    """

    i = 0  # 初始化存储每个块的信息的列表
    text = text.replace("\n", " ").replace(
        "，", " ").replace("。", " ").split()  # 将文本中的换行符替换为空格，并分割成单词列表

    # # 如果预计的API调用次数超出最大限制，抛出错误
    # if (len(text) / chunk_sz) >= max_memory:
    #     raise ValueError(
    #         "Processing is aborted due to high anticipated costs.")

    # 使用tqdm库显示处理进度
    for idx in tqdm(range(0, len(text), chunk_sz)):

        chunk = " ".join(text[idx: idx + chunk_sz])  # 生成当前块的文本

        # print("len(tokenizer.encode(chunk))", len(tokenizer.encode(chunk)))

        # 如果当前块的编码长度超过限制，则跳过
        # if len(tokenizer.encode(chunk)) > chunk_sz * 3:
        #     print("Skipped an uninformative chunk.")
        #     continue

        attempts = 0  # 初始化尝试次数
        while True:
            try:
                # 生成当前文本块的摘要和嵌入向量
                embd = get_embedding(chunk)

                doc1 = Document(id=i,
                                content=chunk,
                                embedding=embd,
                                )

                # print(f"doc_id: {i}, content: {chunk}")

                document_store.write_documents([doc1])
                i += 1
                # time.sleep(3)  # 暂停3秒，以遵守API调用频率限制
                break  # 成功后退出循环

            except Exception as e:
                attempts += 1
                if attempts >= 3:  # 如果尝试次数超过3次，则抛出异常
                    raise Exception(f"{str(e)}")
                time.sleep(3)  # 等待3秒后重试


def retrieve(query):
    """
    基于给定的查询嵌入向量，从一组文本和摘要中检索最相关的三个文本。

    args:
    - q_embd: 查询的嵌入向量，numpy数组类型。
    - info: 包含文本和摘要嵌入向量的信息列表，每个元素都是一个字典。

    returns:
    - 返回最相关的三个文本的索引列表。
    """

    MAX_TOP_N = 3

    print("开始retrieve")
    print("问题：", query)

    locker.acquire()
    documents = doc_retriever.retrieve(query,
                                       top_k=MAX_TOP_N,
                                       document_store=document_store)

    print(f'Got search_doc_embedding_data count: {len(documents)}')
    for doc in documents:
        print(f"doc_id: {doc.id}, content: {doc.content}, score: {doc.score}")
    locker.release()

    return documents


# prompt

def get_question():
    q = input("Enter your question: ")
    return q


def get_qa_content(q, retrieved_doc):

    retrieved_text = [doc.content for doc in retrieved_doc]

    content = "After reading some relevant passage fragments from the same document, please respond to the following query. Note that there may be typographical errors in the passages due to the text being fetched from a PDF file or web page."

    content += "\nQuery: " + q

    for i in range(len(retrieved_text)):
        content += "\nPassage " + str(i + 1) + ": " + retrieved_text[i]

    content += "\nAvoid explicitly using terms such as 'passage 1, 2 or 3' in your answer as the questioner may not know how the fragments are retrieved. Please use the same language as in the query to respond."

    return content

# chat


# def chatGPT_api(messages):
#     # 使用OpenAI的API创建一个聊天机器人的完成请求
#     # args:
#     #   messages: 包含对话历史的列表，每个元素都是一个包含“role”和“content”键的字典
#     # returns:
#     #   返回由模型生成的消息字符串

#     completion = openai.ChatCompletion.create(
#         model='gpt-3.5-turbo',  # 指定使用的模型版本为gpt-3.5-turbo
#         messages=messages,  # 传入对话历史信息
#         temperature=1,  # 设定生成温度，1表示较高的创造性
#         top_p=0.95,  # top_p用于控制生成的多样性，0.95意味着仅考虑累积概率达到95%的最可能的词
#         # max_tokens=2000, # 最大生成的令牌数，这里被注释掉了，默认值或其他设置应在调用前确定
#         frequency_penalty=0.0,  # 频率惩罚设为0，不对重复内容进行惩罚
#         presence_penalty=0.0  # 存在惩罚设为0，不对新内容的出现进行惩罚
#     )
#     return completion.choices[0].message  # 返回模型生成的第一条消息文本


def generate_answer(q, retrieved_documents):
    while True:

        content = get_qa_content(q, retrieved_documents)
        if len(tokenizer.encode(content)) > 3800:
            retrieved_indices = retrieved_indices[:-1]
            print("Contemplating...")
            if not retrieved_indices:
                raise ValueError("Failed to respond.")
        else:
            break
    # messages = [
    #     {"role": "user", "content": content}
    # ]

    # print(content)

    response = requests.post(url, headers=headers,
                             data=makedata(content), stream=True)

    # for line in response.iter_lines():
    #     if line:
    #         text = line.decode("utf-8")  # 将字节流解码为文本
    #         print(text)  # 打印每行文本数据

    res = ""  # 初始化空字符串
    for line in response.iter_lines():
        if line:
            text = line.decode("utf-8")  # 将字节流解码为文本
            res += text + " "  # 将文本添加到结果字符串中，并添加空格以分隔文本
            
    print("res:",res)
    return res
    # return messages


def answer(q):

    retrieved_documents = retrieve(q)
    answer = generate_answer(q, retrieved_documents)
    return answer


def chat(text_path):

    store_info(get_text(text_path))

    while True:
        q = get_question()
        if len(tokenizer.encode(q)) > 200:
            raise ValueError("Input query is too long!")
        attempts = 0
        while True:
            try:
                response = answer(q)
                print()
                print(f"{bcolors.OKGREEN}{response}{bcolors.ENDC}")
                print()
                time.sleep(3)  # up to 20 api calls per min
                break
            except Exception as e:
                attempts += 1
                if attempts >= 3:
                    raise Exception(f"{str(e)}")
                time.sleep(3)
